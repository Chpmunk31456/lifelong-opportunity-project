#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from pypdf import PdfReader
import hashlib, re, difflib

root = Path(__file__).resolve().parents[1]
docx_path = root / '07-customer-service-specialist/english/docx/Lifelong_Opportunity_Customer_Service_Specialist_Guide_English_v1.0.docx'
pdf_path = root / '07-customer-service-specialist/english/pdf/Lifelong_Opportunity_Customer_Service_Specialist_Guide_English_v1.0.pdf'
out = root / 'project/revision-2026/guide-07/qa/evidence'
out.mkdir(parents=True, exist_ok=True)

def sha256(path): return hashlib.sha256(path.read_bytes()).hexdigest()
def clean_lines(text): return '\n'.join(x.rstrip() for x in text.splitlines()).rstrip() + '\n'

doc = Document(docx_path)
blocks=[]; headings=[]
for p in doc.paragraphs:
    t=' '.join(p.text.split())
    if t:
        blocks.append(t)
        if 'heading' in (p.style.name or '').lower(): headings.append(t)
for table in doc.tables:
    for row in table.rows:
        vals=[' '.join(c.text.split()) for c in row.cells]
        vals=[v for v in vals if v]
        if vals: blocks.append(' | '.join(vals))
doc_text='\n'.join(blocks)

reader=PdfReader(str(pdf_path))
pages=[]
for i,page in enumerate(reader.pages,1):
    txt=(page.extract_text() or '').replace('\x00','').strip()
    pages.append(f'--- PAGE {i} ---\n{txt}')
pdf_text='\n'.join(pages)
(out/'guide07_legacy_docx_extract.txt').write_text(clean_lines(doc_text),encoding='utf-8')
(out/'guide07_legacy_pdf_extract.txt').write_text(clean_lines(pdf_text),encoding='utf-8')

def norm(text):
    text=text.lower().replace('–','-').replace('—','-').replace('’',"'")
    text=re.sub(r'--- page \d+ ---',' ',text)
    text=re.sub(r'\s+',' ',text)
    return re.sub(r"[^a-z0-9$%+./#' -]",'',text).strip()
nd,npdf=norm(doc_text),norm(pdf_text)
sequence=difflib.SequenceMatcher(None,nd,npdf).ratio()
pattern=r"[a-z0-9][a-z0-9$%+./#'-]*"
dt=set(re.findall(pattern,nd)); pt=set(re.findall(pattern,npdf)); union=dt|pt
jaccard=len(dt&pt)/len(union) if union else 1.0
fact=re.compile(r'(?<![A-Za-z])(?:\$?\d[\d,.]*%?|BLS|WIOA|SENA|NOC)',re.I)
df=set(v.lower() for v in fact.findall(doc_text)); pf=set(v.lower() for v in fact.findall(pdf_text))
only_doc=sorted(df-pf); only_pdf=sorted(pf-df)
if sequence>=.90 and jaccard>=.90 and not only_doc and not only_pdf: classification='HIGH_EQUIVALENCE'
elif sequence>=.80 and jaccard>=.85 and len(only_doc)+len(only_pdf)<=4: classification='PROBABLE_EQUIVALENCE_REVIEW_REQUIRED'
else: classification='RECONCILE'
report=[
'# Guide 07 — Legacy English Extraction and Reconciliation 02','',
'Date: 2026-08-08','Branch: `revision/guide-00-100-2026`','Guide: 07 — Customer Service Specialist','',
'## Gate purpose','',
'Deterministically extract the legacy English DOCX and searchable PDF, preserve both text extractions as audit evidence, and measure whether one artifact can safely be treated as the sole source for the 2026 reconstruction. This is a source-reconciliation control, not factual revalidation, publication approval, independent human review, certification, accreditation, or accessibility certification.','',
'## Artifact fingerprints','',
'- DOCX Git blob SHA: `3f78d7f6b78769372403230b83e0e75ef3e13535`',
'- PDF Git blob SHA: `3a9faaca6757997b511c5cb73e21eb48b5ecc0e1`',
f'- DOCX SHA-256: `{sha256(docx_path)}`',f'- PDF SHA-256: `{sha256(pdf_path)}`',
f'- DOCX extracted non-empty blocks: {len(blocks)}',f'- PDF pages: {len(reader.pages)}',
f'- DOCX extracted characters: {len(doc_text):,}',f'- PDF extracted characters: {len(pdf_text):,}','',
'## Deterministic comparison','',f'- Normalized character-sequence similarity: **{sequence:.4f}**',f'- Normalized unique-token Jaccard similarity: **{jaccard:.4f}**',
f"- Material token/fact set only in DOCX: `{', '.join(only_doc) or 'none detected'}`",f"- Material token/fact set only in PDF: `{', '.join(only_pdf) or 'none detected'}`",f'- Automated classification: **{classification}**','',
'## DOCX heading inventory','']
report += ['- '+h for h in headings[:100]]
report += ['','## Evidence files','','- `guide07_legacy_docx_extract.txt`','- `guide07_legacy_pdf_extract.txt`','','## Controlled decision','','**HOLD for substantive reconciliation.** Automated similarity does not by itself prove substantive equivalence. Review substantive sections and record the composite-source decision before drafting the revised English master.','']
(root/'project/revision-2026/guide-07/qa/GUIDE_07_LEGACY_ENGLISH_EXTRACTION_RECONCILIATION_02.md').write_text(clean_lines('\n'.join(report)),encoding='utf-8')
print(f'classification={classification}')
print(f'sequence_similarity={sequence:.4f}')
print(f'token_jaccard={jaccard:.4f}')
